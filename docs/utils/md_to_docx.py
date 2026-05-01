#!/usr/bin/env python3
"""Convert Markdown files to branded CaliLean DOCX documents.

Uses brand typography and colors from docs/brand/brand-book/CaliLean_Brand_Elements.html
and docs/utils/brand_constants.py. Logo placed top-right in header.

Usage:
    python3 docs/utils/md_to_docx.py docs/strategy/identity-brief.md
    python3 docs/utils/md_to_docx.py --all          # Convert all mapped MDs
    python3 docs/utils/md_to_docx.py --all --clean   # Remove generated DOCX files
"""

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent

# ── Brand constants ──────────────────────────────────────────────────
INK = RGBColor(0x11, 0x11, 0x11)
PACIFIC = RGBColor(0x70, 0x90, 0xAB)
FOG = RGBColor(0x9C, 0xA3, 0xA8)
SAND = RGBColor(0xF0, 0xF0, 0xF0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ALERT = RGBColor(0xA2, 0x3B, 0x2A)

DISPLAY_FONT = "Instrument Serif"
BODY_FONT = "Plus Jakarta Sans"
MONO_FONT = "JetBrains Mono"

LOGO_PATH = PROJECT_ROOT / "docs" / "brand" / "assets" / "logo" / "master" / "calilean-logo-black.png"

# Output directory for generated DOCX files
BUILD_DIR = PROJECT_ROOT / "docs" / "build"

# Which MD directories to convert (mirrors SYNC_MAP in sync_drive.py)
MD_SOURCES = [
    PROJECT_ROOT / "docs" / "strategy",
    PROJECT_ROOT / "docs" / "brand" / "packaging" / "ideation",
    PROJECT_ROOT / "docs" / "brand" / "assets" / "imagery",
]


def _hex_str(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


# ── Template: create a blank branded document ────────────────────────

def create_branded_doc() -> Document:
    """Create a new Document with CaliLean brand styles and header logo."""
    doc = Document()

    # Page setup — US Letter, 1-inch margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Header with logo top-right ──
    header = section.header
    header.is_linked_to_previous = False

    # Remove default empty paragraph that adds vertical space
    for p in header.paragraphs:
        p._p.getparent().remove(p._p)

    hp = header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(4)
    if LOGO_PATH.exists():
        run = hp.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.6))

    # Thin line under header
    pPr = hp._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="{_hex_str(SAND)}"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # ── Footer ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(fp, "hello@calilean.com  \u00b7  calilean.com", BODY_FONT, Pt(8), FOG)

    # ── Define styles ──
    _define_styles(doc)

    return doc


def _define_styles(doc: Document):
    """Register branded paragraph and character styles."""
    styles = doc.styles

    # Title style
    if "CL Title" not in [s.name for s in styles]:
        title_style = styles.add_style("CL Title", 1)  # 1 = paragraph
        title_style.font.name = DISPLAY_FONT
        title_style.font.size = Pt(22)
        title_style.font.color.rgb = INK
        title_style.paragraph_format.space_before = Pt(0)
        title_style.paragraph_format.space_after = Pt(6)

    # Subtitle / metadata
    if "CL Meta" not in [s.name for s in styles]:
        meta_style = styles.add_style("CL Meta", 1)
        meta_style.font.name = BODY_FONT
        meta_style.font.size = Pt(9)
        meta_style.font.color.rgb = FOG
        meta_style.font.italic = True
        meta_style.paragraph_format.space_before = Pt(0)
        meta_style.paragraph_format.space_after = Pt(2)

    # H2
    if "CL H2" not in [s.name for s in styles]:
        h2 = styles.add_style("CL H2", 1)
        h2.font.name = BODY_FONT
        h2.font.size = Pt(16)
        h2.font.bold = True
        h2.font.color.rgb = INK
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(6)

    # H3
    if "CL H3" not in [s.name for s in styles]:
        h3 = styles.add_style("CL H3", 1)
        h3.font.name = BODY_FONT
        h3.font.size = Pt(13)
        h3.font.bold = True
        h3.font.color.rgb = INK
        h3.paragraph_format.space_before = Pt(14)
        h3.paragraph_format.space_after = Pt(4)

    # H4
    if "CL H4" not in [s.name for s in styles]:
        h4 = styles.add_style("CL H4", 1)
        h4.font.name = BODY_FONT
        h4.font.size = Pt(11)
        h4.font.bold = True
        h4.font.color.rgb = PACIFIC
        h4.paragraph_format.space_before = Pt(10)
        h4.paragraph_format.space_after = Pt(3)

    # Body
    if "CL Body" not in [s.name for s in styles]:
        body = styles.add_style("CL Body", 1)
        body.font.name = BODY_FONT
        body.font.size = Pt(10)
        body.font.color.rgb = INK
        body.paragraph_format.space_before = Pt(0)
        body.paragraph_format.space_after = Pt(6)
        body.paragraph_format.line_spacing = 1.35

    # List item
    if "CL List" not in [s.name for s in styles]:
        li = styles.add_style("CL List", 1)
        li.font.name = BODY_FONT
        li.font.size = Pt(10)
        li.font.color.rgb = INK
        li.paragraph_format.space_before = Pt(1)
        li.paragraph_format.space_after = Pt(1)
        li.paragraph_format.left_indent = Inches(0.35)
        li.paragraph_format.line_spacing = 1.3

    # Blockquote
    if "CL Quote" not in [s.name for s in styles]:
        bq = styles.add_style("CL Quote", 1)
        bq.font.name = BODY_FONT
        bq.font.size = Pt(9)
        bq.font.color.rgb = FOG
        bq.font.italic = True
        bq.paragraph_format.space_before = Pt(2)
        bq.paragraph_format.space_after = Pt(2)
        bq.paragraph_format.left_indent = Inches(0.3)

    # Code block
    if "CL Code" not in [s.name for s in styles]:
        code = styles.add_style("CL Code", 1)
        code.font.name = MONO_FONT
        code.font.size = Pt(8.5)
        code.font.color.rgb = PACIFIC
        code.paragraph_format.space_before = Pt(4)
        code.paragraph_format.space_after = Pt(4)
        code.paragraph_format.left_indent = Inches(0.3)


# ── Inline formatting helpers ────────────────────────────────────────

def _add_run(paragraph, text, font_name=None, size=None, color=None,
             bold=False, italic=False, mono=False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    if mono:
        run.font.name = MONO_FONT
        run.font.size = size or Pt(9)
        run.font.color.rgb = PACIFIC
    else:
        if font_name:
            run.font.name = font_name
        if size:
            run.font.size = size
        if color:
            run.font.color.rgb = color
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return run


def _add_inline_formatted(paragraph, text, style_name="CL Body"):
    """Parse inline markdown (bold, italic, code) and add runs to paragraph."""
    # Split by inline patterns: **bold**, *italic*, `code`
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            _add_run(paragraph, part[2:-2], BODY_FONT, Pt(10), INK, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            _add_run(paragraph, part[1:-1], BODY_FONT, Pt(10), INK, italic=True)
        elif part.startswith('`') and part.endswith('`'):
            _add_run(paragraph, part[1:-1], mono=True, size=Pt(9))
        else:
            _add_run(paragraph, part, BODY_FONT, Pt(10), INK)


# ── Markdown parsing and rendering ───────────────────────────────────

def parse_md_to_docx(md_text: str, doc: Document):
    """Parse markdown text and fill content into the branded DOCX template."""
    lines = md_text.split('\n')
    i = 0
    title_found = False
    h2_count = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # ── Code blocks ──
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block — flush
                p = doc.add_paragraph(style="CL Code")
                # Add shading behind code block
                pPr = p._p.get_or_add_pPr()
                shd = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{_hex_str(SAND)}" w:val="clear"/>'
                )
                pPr.append(shd)
                _add_run(p, '\n'.join(code_lines), MONO_FONT, Pt(8.5), PACIFIC)
                code_lines = []
                in_code_block = False
            else:
                # Flush any pending table
                if in_table:
                    _render_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # ── Empty lines ──
        if not stripped:
            if in_table:
                _render_table(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        # ── Horizontal rule ──
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            if in_table:
                _render_table(doc, table_rows)
                table_rows = []
                in_table = False
            # Add a thin line
            p = doc.add_paragraph()
            p.space_before = Pt(8)
            p.space_after = Pt(8)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{_hex_str(SAND)}"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # ── Tables ──
        if '|' in stripped and stripped.startswith('|'):
            # Skip separator rows like |---|---|
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if all(re.match(r'^[-:]+$', c) for c in cells if c):
                i += 1
                continue
            table_rows.append(cells)
            in_table = True
            i += 1
            continue

        # Flush pending table if we hit a non-table line
        if in_table:
            _render_table(doc, table_rows)
            table_rows = []
            in_table = False

        # ── Headers ──
        h_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()

            if level == 1 and not title_found:
                p = doc.add_paragraph(style="CL Title")
                _add_run(p, text, DISPLAY_FONT, Pt(22), INK)
                title_found = True
            elif level == 1:
                p = doc.add_paragraph(style="CL H2")
                p.paragraph_format.page_break_before = True
                _add_inline_formatted(p, text)
            elif level == 2:
                h2_count += 1
                p = doc.add_paragraph(style="CL H2")
                # Page break before H2 sections (skip the first one after title)
                if h2_count > 1:
                    p.paragraph_format.page_break_before = True
                _add_inline_formatted(p, text)
            elif level == 3:
                p = doc.add_paragraph(style="CL H3")
                _add_inline_formatted(p, text)
            elif level == 4:
                p = doc.add_paragraph(style="CL H4")
                _add_inline_formatted(p, text)

            i += 1
            continue

        # ── Blockquotes ──
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('>').strip()
            p = doc.add_paragraph(style="CL Quote")
            # Add left border accent
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="12" w:space="4" w:color="{_hex_str(PACIFIC)}"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            _add_inline_formatted(p, quote_text, "CL Quote")
            i += 1
            continue

        # ── Ordered lists ──
        ol_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if ol_match:
            num = ol_match.group(1)
            text = ol_match.group(2)
            p = doc.add_paragraph(style="CL List")
            _add_run(p, f"{num}. ", BODY_FONT, Pt(10), FOG, bold=True)
            _add_inline_formatted(p, text)
            i += 1
            continue

        # ── Unordered lists ──
        ul_match = re.match(r'^[-*+]\s+(.*)', stripped)
        if ul_match:
            text = ul_match.group(1)
            p = doc.add_paragraph(style="CL List")
            _add_run(p, "\u2022  ", BODY_FONT, Pt(10), PACIFIC)
            _add_inline_formatted(p, text)
            i += 1
            continue

        # ── Body text ──
        p = doc.add_paragraph(style="CL Body")
        _add_inline_formatted(p, stripped)
        i += 1

    # Flush remaining table
    if in_table:
        _render_table(doc, table_rows)


def _render_table(doc: Document, rows: list):
    """Render a markdown table as a styled DOCX table."""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style: light grid
    table.style = "Table Grid"

    for row_idx, cells in enumerate(rows):
        for col_idx, cell_text in enumerate(cells):
            if col_idx >= num_cols:
                break
            cell = table.cell(row_idx, col_idx)

            # Clear default paragraph
            p = cell.paragraphs[0]
            p.clear()

            if row_idx == 0:
                # Header row: bold, sand background
                _add_run(p, cell_text.strip(), BODY_FONT, Pt(9), INK, bold=True)
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{_hex_str(SAND)}" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)
            else:
                _add_inline_formatted(p, cell_text.strip())
                # Set font size for table body
                for run in p.runs:
                    run.font.size = Pt(9)

    # Set table borders to subtle color
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{_hex_str(FOG)}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{_hex_str(FOG)}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{_hex_str(FOG)}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{_hex_str(FOG)}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{_hex_str(FOG)}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{_hex_str(FOG)}"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


# ── File conversion ──────────────────────────────────────────────────

def convert_md_file(md_path: Path, out_path: Path = None) -> Path:
    """Convert a single MD file to a branded DOCX. Returns output path."""
    if out_path is None:
        out_path = BUILD_DIR / md_path.relative_to(PROJECT_ROOT).with_suffix('.docx')

    out_path.parent.mkdir(parents=True, exist_ok=True)

    md_text = md_path.read_text(encoding='utf-8')
    doc = create_branded_doc()
    parse_md_to_docx(md_text, doc)
    doc.save(str(out_path))

    print(f"  Built: {md_path.relative_to(PROJECT_ROOT)} -> {out_path.relative_to(PROJECT_ROOT)}")
    return out_path


def find_all_md_files() -> list:
    """Find all MD files in source directories (excluding archive/)."""
    md_files = []
    for src_dir in MD_SOURCES:
        if not src_dir.exists():
            continue
        for f in sorted(src_dir.rglob('*.md')):
            # Skip archive directories and log files
            if 'archive' in f.relative_to(PROJECT_ROOT).parts:
                continue
            if f.name in ('log.md',):
                continue
            md_files.append(f)
    return md_files


def build_all():
    """Convert all mapped MD files to DOCX."""
    md_files = find_all_md_files()
    if not md_files:
        print("No MD files found to convert.")
        return []

    print(f"Converting {len(md_files)} MD file(s) to branded DOCX...\n")
    outputs = []
    for md_file in md_files:
        out = convert_md_file(md_file)
        outputs.append(out)
    print(f"\nDone. {len(outputs)} DOCX file(s) in {BUILD_DIR.relative_to(PROJECT_ROOT)}/")
    return outputs


def clean():
    """Remove the build directory."""
    import shutil
    if BUILD_DIR.exists():
        shutil.rmtree(BUILD_DIR)
        print(f"Cleaned {BUILD_DIR.relative_to(PROJECT_ROOT)}/")
    else:
        print("Nothing to clean.")


# ── CLI ──────────────────────────────────────────────────────────────

def main():
    args = sys.argv[1:]

    if not args:
        print("Usage:")
        print("  python3 docs/utils/md_to_docx.py <file.md> [output.docx]")
        print("  python3 docs/utils/md_to_docx.py --all")
        print("  python3 docs/utils/md_to_docx.py --all --clean")
        print("  python3 docs/utils/md_to_docx.py --clean")
        sys.exit(1)

    if '--clean' in args and '--all' not in args:
        clean()
        return

    if '--all' in args:
        if '--clean' in args:
            clean()
        build_all()
        return

    # Single file conversion
    md_path = PROJECT_ROOT / args[0]
    if not md_path.exists():
        md_path = Path(args[0])
    if not md_path.exists():
        print(f"ERROR: File not found: {args[0]}")
        sys.exit(1)

    out_path = None
    if len(args) > 1 and not args[1].startswith('--'):
        out_path = Path(args[1])
        if not out_path.is_absolute():
            out_path = PROJECT_ROOT / out_path

    convert_md_file(md_path, out_path)


if __name__ == '__main__':
    main()
